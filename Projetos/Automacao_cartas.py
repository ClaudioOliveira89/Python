import openpyxl
from docx import Document 
from docx.shared import Pt

def gerar_cartas():

    wb = openpyxl.load_workbook('/caminho/.xlsx')
    sheet = wb['Planilha1']

    modelo_carta_paht = '/caminho/docx'

    lin =1

    while sheet.cell(row=lin, column=1).value is not None:
        try:
        #para abrir o documento de modelo.
        doc = Document(modelo_carta_path)

        for col in range(1, 16):
            placeholder = sheet.cell(row=1, column=col).value
            replace_text = sheet.cell(row=lin, column=col).value

        for paragraph in doc.paragraphs:
            if placeholder in paragraph.text:
                inline = paragraph.runs
                for item in inline:
                    if placeholder in item.text:
                        item.text = item.text.replace(placeholder, str(replace_text))
                        item.font.size =Pt(12) #aqui pode ajustar a fonte.
        # Substitui os placeholders em tabelas.
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        if placeholder in paragraph.text:
                            inline = paragraph.runs
                            for item in inline:
                                if placeholder in item.text:
                                    item.text = item.text.replace(placeholder, str(replace_text) )
                                    item.font.size = Pt(12) 

#salvando o documento com base nos valores da célula
novo_doc_path =f'path/to/save/cartas_{sheet.cell(row=lin, column=1).value}.docx'
doc.save(novo_doc_path)

# incrementa a linha 
lin +=1

except Exception as e:
    print(f"Erro ao processar a linha {lin}: {e}")
    lin += 1

# para chamar a função
gerar_cartas()